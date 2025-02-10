import docx
import ast
import openai

def generate_documentation(parsed_data, output_file="documentation.docx"):
    """
    Generates documentation from parsed data by utilizing LLM to understand the code and saves it as a Word document.
    """
    doc = docx.Document()
    doc.add_heading('Project Documentation', level=1)
    for data in parsed_data:
        file_name = data["file"]
        for node in ast.walk(data["tree"]):
            if isinstance(node, (ast.FunctionDef, ast.ClassDef)):
                # Use LLM to understand the code and generate descriptions
                code_snippet = ast.unparse(node)
                llm_response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "user", "content": f"Analyze the following code and provide a detailed description of its purpose and functionality:\n{code_snippet}"}
                    ],
                    max_tokens=200
                )
                description = llm_response.choices[0].message['content'].strip()
                # Add heading and description to the document
                if isinstance(node, ast.FunctionDef):
                    doc.add_heading(f'Function: {node.name} (in {file_name})', level=3)
                elif isinstance(node, ast.ClassDef):
                    doc.add_heading(f'Class: {node.name} (in {file_name})', level=2)
                doc.add_paragraph(description)
    doc.save(output_file)
    print(f"Documentation saved as a Word document to {output_file}")